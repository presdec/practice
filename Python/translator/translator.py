try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
from googletrans import Translator
from easygui import choicebox, textbox, fileopenbox
from docx import Document
import os
import pyperclip

"""
Module that provides a destination language choice. Further languages can be
added from the following list under choices:
https://py-googletrans.readthedocs.io/en/latest/#googletrans-languages
"""

source_lang = "en"
choices = ["fr", "el"]
desti_lang = choicebox(
    "What is the Destination Language?", "translator", choices=choices
)


"""
Module that extracts text from MS XML Word document (.docx).
(Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
"""

WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
PARA = WORD_NAMESPACE + "p"
TEXT = WORD_NAMESPACE + "t"


def get_docx_text(filename):
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    document = zipfile.ZipFile(filename)
    xml_content = document.read("word/document.xml")
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [node.text for node in paragraph.getiterator(TEXT) if node.text]
        if texts:
            paragraphs.append("".join(texts))

    return "\n\n".join(paragraphs)


def splitIntoSentences(text):
    text = text.split("\n")
    return text


def translate(text, source, dest):
    translator = Translator()
    if text is None:
        print("Invalid Text")
    elif text != "":
        translation = translator.translate(text, src=source, dest=dest).text
        pyperclip.copy(translation)
    return translation


def main(filename, source_lang, desti_lang):
    print("Translating:  " + os.path.basename(filename))
    count = 0
    clean = get_docx_text(filename)
    para = splitIntoSentences(clean)
    final = []
    for p in para:
        p = p.split(".")
        tmp = []
        for sentence in p:
            count += 1
            if sentence is None:
                print("Invalid Text")
            elif sentence != "" and not None:
                print(str(count) + ":" + str(sentence))
                translatedText = translate(sentence, source_lang, desti_lang)
                message = (
                    "Original: \n"
                    + str(sentence)
                    + "\nGoogleTrans:\n"
                    + str(translatedText)
                )
                fd_v = textbox(
                    msg=str(message), title="translation", text=translatedText
                )
                tmp.append(fd_v)
        final.append(". ".join(tmp))
    final = "\n".join(final)
    f = open("translation.txt", "w")
    f.writelines(final)
    print("please check translation.txt")
    f.close()
    document = Document()
    document.add_heading("Translation", 0)
    p = document.add_paragraph(final)
    document.add_page_break()
    document.save("translated_" + os.path.basename(filename))


if __name__ == "__main__":
    filename = fileopenbox()
    main(filename, source_lang, desti_lang)
